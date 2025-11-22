"""Unit tests for HybridCVAnalyzerAdapter."""

from __future__ import annotations

from pathlib import Path
from typing import Any
from uuid import uuid4

import pytest
from docx import Document

from src.adapters.cv_processing.hybrid_cv_analyzer_adapter import (
    HybridCVAnalyzerAdapter,
)
from src.adapters.cv_processing.llm_fallback_extractor import (
    LLMFallbackExtractorProtocol,
)
from src.domain.models.cv_skill import CVSkill, ProficiencyLevel


class StubRuleExtractor:
    def __init__(self, response: dict[str, Any]) -> None:
        self.response = response
        self.calls = 0

    def extract(self, cv_text: str) -> dict[str, Any]:
        self.calls += 1
        return self.response


class StubNERExtractor:
    def __init__(self, response: dict[str, Any]) -> None:
        self.response = response
        self.calls = 0

    def extract(self, cv_text: str) -> dict[str, Any]:
        self.calls += 1
        return self.response


class StubLLMFallback(LLMFallbackExtractorProtocol):
    def __init__(self) -> None:
        self.called = False

    async def fill_gaps(
        self,
        merged_results: dict[str, Any],
        cv_text: str,
    ) -> dict[str, Any]:
        self.called = True
        merged_results["llm_enriched"] = True
        return merged_results


def build_rule_response() -> dict[str, Any]:
    return {
        "emails": ["jane@doe.dev"],
        "phones": ["+1-202-555-0101"],
        "dates": ["2020"],
        "sections": {"EXPERIENCE": (1, "Experience")},
        "urls": ["https://example.com"],
    }


def build_ner_response(overall: float = 0.9) -> dict[str, Any]:
    return {
        "name": "Jane Doe",
        "companies": ["Example Corp"],
        "locations": ["San Francisco"],
        "dates": ["2020"],
        "skills": [{"skill": "Python", "category": "programming"}],  # Return dict instead of ExtractedSkill
        "experience_years": 5.0,
        "confidence": {
            "fields": {
                "name": 0.9,
                "companies": 0.85,
                "locations": 0.8,
                "skills": 0.9,
            },
            "overall": overall,
        },
        "language": "en",
    }


@pytest.mark.asyncio
async def test_analyze_cv_high_confidence_skips_llm(tmp_path: Path) -> None:
    cv_path = tmp_path / "cv.txt"
    cv_path.write_text("Sample CV", encoding="utf-8")

    llm_stub = StubLLMFallback()
    adapter = HybridCVAnalyzerAdapter(
        rule_extractor=StubRuleExtractor(build_rule_response()),
        ner_extractor=StubNERExtractor(build_ner_response()),
        llm_fallback=llm_stub,
        confidence_threshold=0.7,
    )

    analysis = await adapter.analyze_cv(str(cv_path), uuid4())

    assert analysis.metadata["extraction_method"] == "hybrid"
    assert analysis.metadata["confidence"]["email"] >= 0.9
    assert llm_stub.called is False
    assert analysis.skills[0].skill == "Python"


@pytest.mark.asyncio
async def test_analyze_cv_low_confidence_triggers_llm(tmp_path: Path) -> None:
    cv_path = tmp_path / "cv.txt"
    cv_path.write_text("Sample CV missing emails", encoding="utf-8")

    rule_response = build_rule_response()
    rule_response["emails"] = []
    llm_stub = StubLLMFallback()

    adapter = HybridCVAnalyzerAdapter(
        rule_extractor=StubRuleExtractor(rule_response),
        ner_extractor=StubNERExtractor(build_ner_response(overall=0.2)),
        llm_fallback=llm_stub,
        confidence_threshold=0.7,
    )

    analysis = await adapter.analyze_cv(str(cv_path), uuid4())

    assert llm_stub.called is True
    assert analysis.metadata["confidence"]["email"] == 0.0
    assert analysis.metadata.get("language") == "en"


@pytest.mark.asyncio
async def test_generate_candidate_from_summary_uses_rules_and_ner(tmp_path: Path) -> None:
    adapter = HybridCVAnalyzerAdapter(
        rule_extractor=StubRuleExtractor(
            {
                "emails": ["candidate@example.com"],
                "phones": [],
                "dates": [],
                "sections": {},
                "urls": [],
            }
        ),
        ner_extractor=StubNERExtractor(build_ner_response()),
    )

    candidate = await adapter.generate_candidate_from_summary(
        "Jane Doe can be reached at candidate@example.com",
        str(tmp_path / "summary.txt"),
        uuid4(),
    )

    assert candidate.name == "Jane Doe"
    assert candidate.email == "candidate@example.com"


def test_read_docx_supported(tmp_path: Path) -> None:
    doc_path = tmp_path / "cv.docx"
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Senior Engineer")
    document.save(doc_path)

    adapter = HybridCVAnalyzerAdapter(
        rule_extractor=StubRuleExtractor(build_rule_response()),
        ner_extractor=StubNERExtractor(build_ner_response()),
    )

    content = adapter._read_cv_text(str(doc_path))

    assert "Jane Doe" in content
    assert "Senior Engineer" in content

