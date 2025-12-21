"""Unit tests for LLMCVAnalyzerAdapter."""

from __future__ import annotations

from pathlib import Path
from typing import Any
from uuid import uuid4

import pytest
from docx import Document

from src.infrastructure.adapters.cv_processing.llm_cv_analyzer_adapter import LLMCVAnalyzerAdapter
from src.infrastructure.adapters.llm.cv_skill_extraction_models import (
    CVSkillExtractionOutput,
    SkillOutput,
    ProficiencyLevelOutput,
)


class MockLLMPort:
    """Mock LLM port for testing."""

    def __init__(self, extraction_result: CVSkillExtractionOutput | None = None) -> None:
        self.extraction_result = extraction_result or CVSkillExtractionOutput(
            skills=[
                SkillOutput(
                    skill_name="Python",
                    proficiency_level=ProficiencyLevelOutput.ADVANCED,
                    years_of_experience=5.0,
                    is_primary=True,
                ),
                SkillOutput(
                    skill_name="FastAPI",
                    proficiency_level=ProficiencyLevelOutput.INTERMEDIATE,
                    years_of_experience=2.0,
                    is_primary=False,
                ),
            ],
            summary="Senior Python developer with 5 years experience.",
        )
        self.calls: list[dict[str, Any]] = []

    async def analyze_cv_for_skills(
        self, cv_text: str, context: dict[str, Any] | None = None
    ) -> CVSkillExtractionOutput:
        self.calls.append({"cv_text": cv_text, "context": context})
        return self.extraction_result


class FailingLLMPort:
    """Mock LLM port that always fails."""

    async def analyze_cv_for_skills(
        self, cv_text: str, context: dict[str, Any] | None = None
    ) -> CVSkillExtractionOutput:
        raise RuntimeError("LLM extraction failed")


@pytest.mark.asyncio
async def test_analyze_cv_success(tmp_path: Path) -> None:
    """Test successful CV analysis."""
    cv_path = tmp_path / "cv.txt"
    cv_path.write_text("Senior Python Developer with 5 years experience", encoding="utf-8")
    cv_bytes = cv_path.read_bytes()

    mock_llm = MockLLMPort()
    adapter = LLMCVAnalyzerAdapter(llm_port=mock_llm)

    analysis = await adapter.analyze_cv(cv_bytes, "txt", str(uuid4()))

    assert len(mock_llm.calls) == 1
    assert len(analysis.skills) == 2
    assert analysis.skills[0].skill_name == "Python"
    assert analysis.summary == "Senior Python developer with 5 years experience."


@pytest.mark.asyncio
async def test_analyze_cv_fails_fast_on_llm_error(tmp_path: Path) -> None:
    """Test fail-fast behavior when LLM fails."""
    cv_path = tmp_path / "cv.txt"
    cv_path.write_text("Sample CV", encoding="utf-8")
    cv_bytes = cv_path.read_bytes()

    adapter = LLMCVAnalyzerAdapter(llm_port=FailingLLMPort())

    with pytest.raises(RuntimeError, match="LLM extraction failed"):
        await adapter.analyze_cv(cv_bytes, "txt", str(uuid4()))


def test_constructor_requires_llm_port() -> None:
    """Test that llm_port is required."""
    with pytest.raises(ValueError, match="llm_port is required"):
        LLMCVAnalyzerAdapter(llm_port=None)  # type: ignore[arg-type]


def test_read_docx_from_bytes(tmp_path: Path) -> None:
    """Test DOCX parsing."""
    doc_path = tmp_path / "cv.docx"
    document = Document()
    document.add_paragraph("Jane Doe")
    document.add_paragraph("Senior Engineer")
    document.save(doc_path)
    doc_bytes = doc_path.read_bytes()

    mock_llm = MockLLMPort()
    adapter = LLMCVAnalyzerAdapter(llm_port=mock_llm)

    content = adapter._read_cv_text_from_bytes(doc_bytes, "docx")

    assert "Jane Doe" in content
    assert "Senior Engineer" in content
